from rest_framework import viewsets, mixins
from App.models import Image, Subject, Question, Answer
from App.serializers import QuestionDeleteSerializer, QuestionSerializer, QuestionValidate, QuestionCreateSerializer
from App.commons.response import ResponseReadMany, ResponseReadOne, ResponseCreateOne, ResponseDestroyOne
from App.commons.enum import ReponseEnum
from rest_framework.decorators import action
from docx import Document
#pip install python-docx

class ImportView(
        viewsets.GenericViewSet
):
    queryset = Image.objects.all()

    SUCCESS = "success"
    ERROR = "error"

    @action(detail=False, methods=['post'], serializer_class=QuestionSerializer)
    def read_import_file(self, path):
        doc = Document(path)
        content = []
        examInfor = {
               "subject": False,
               "num_of_question": False,
               "lecturer": False,
                "date": False
        }
        # Read string
        for para in doc.paragraphs:
                text = para.text.strip().lower
                if text == "": continue
                checkValid = validate_infor(text, examInfor)
                if checkValid["status"] == "error":
                       raise Exception(checkValid["message"])
                else: content.append(text)
        for key, value in examInfor.items():
               if value == False: raise Exception("Exam must have {} value", key)
        
        # Read tables
        listQuestion = enumerate(doc.tables)
        if len(listQuestion) <= 0: raise Exception("Exam must have question")
        for table_index, table in listQuestion:
                table_data = []
                question_infor = {
                        "question": False,
                        "answer": False,
                        "correct_answer": False,
                        "mark": False,
                        "unit": False,
                        "mix_choice": False
                }
                for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        checkValid = validate_question(row_data, question_infor)
                        if checkValid["status"] == "error":
                                raise Exception("Question index: "+str(table_index)+checkValid["message"])
                        else: table_data.append(row_data)
                content.append({'table_index': table_index, 'data': table_data})
                for key, value in question_infor.items():
                        if value == False: raise Exception("Question {} must have {} value", (table_index,key))
        
        return content
    
    def validate_infor(text, examInfor):
        result = {"status":"error", "message":""}
        message = ""
        if text.lower().startswith('subject'):
                subjectCode = text.split()[1]
                if subjectCode is None or subjectCode == "": 
                       message="Exam must have subject code"
                examInfor["subject"] = True
        elif text.lower().startswith('number'):
                numOfQuiz = text.split()
                numOfQuiz = numOfQuiz[len(numOfQuiz)-1]
                try:
                       numOfQuiz = int(numOfQuiz)
                       if numOfQuiz<=0:
                                message = "Number of total question must be greater than 0"
                except:
                       message = "Number of total question must be integer"
                examInfor["num_of_question"] = True
        elif text.lower().startswith('lecturer'):
                lecturer = text.split()[1]
                if lecturer is None or lecturer == "": 
                       message="Exam must have lecturer name"
                examInfor["lecturer"] = True
        elif text.lower().startswith('date'):
                day = text.split()[1]
                try:
                       date_object = datetime.strptime(day, '%d-%m-%Y')
                except:
                       message = "Date must be a valid day with format dd-mm-yyyy"
                examInfor["date"] = True
        if message=="": result["status"] = "success"
        else: result["message"] = message
        return result
                      

    def validate_question(row, infor):
        result = {"status":"error", "message":""}
        message = ""
        if row[0].lower().startswith('qn'):
                if row[1] == None or row[1] == "":
                      message = "Question must have content of question"
                infor["question"] = True
        elif row[0].lower().startswith('answer'):
                if row[1] == None or row[1] == "":
                      message = "Question must have correct answer"
                infor["correct_answer"] = True
        elif row[0].lower().startswith('mark'):
                try:
                        mark = float(row[1])
                        if mark <= 0: message = "Mark must be greater than 0"
                        infor["mark"] = True
                except:
                       message = "Mark must be float value"
        elif row[0].lower().startswith('unit'):
                if row[1] == None or row[1] == "":
                      message = "Question must have unit"
                infor["unit"] = True
        elif row[0].lower().startswith('mix'):
                try:
                        choice = bool(row[1])
                        infor["mix_choice"] = True
                except: message = "Question must have mix choice"
        elif row[0].lower().startswith('a') or row[0].lower().startswith('b'):
                if row[1] == None or row[1] == "":
                      message = "Question must have answer"
                infor["answer"] = True
        if message=="": result["status"] = "success"
        else: result["message"] = message
        return result
        
                
